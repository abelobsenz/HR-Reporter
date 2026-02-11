from __future__ import annotations

from concurrent.futures import ThreadPoolExecutor, as_completed
import csv
from datetime import datetime, timezone
import hashlib
import logging
import math
import os
import re
import threading
import time
from dataclasses import dataclass, replace
from pathlib import Path
from typing import Dict, List, Sequence, Tuple
from urllib.parse import urldefrag, urljoin, urlparse

from app.models import RawDocument
from app.ingest.text_normalize import normalize_text


SUPPORTED_EXTENSIONS = {".txt", ".md", ".pdf", ".docx", ".csv"}
DIRECTORY_LINK_KEYWORDS = [
    "policy",
    "handbook",
    "time-off",
    "leave",
    "onboarding",
    "hiring",
    "compensation",
    "benefits",
    "performance",
    "manager",
    "relations",
    "harassment",
    "workday",
]
HR_URL_PATH_TOKENS = [
    "people",
    "policy",
    "policies",
    "handbook",
    "hiring",
    "talent",
    "onboarding",
    "performance",
    "manager",
    "benefits",
    "leave",
    "overtime",
    "harassment",
    "employment-law",
    "risk-management",
    "dispute-resolution",
    "publiccompanyresources",
    "legal/employment",
]
HR_CONTENT_TOKENS = [
    "employee handbook",
    "anti-harassment",
    "harassment policy",
    "leave policy",
    "overtime",
    "manager training",
    "performance review",
    "goal framework",
    "pay band",
    "leveling",
    "hris",
    "benefits eligibility",
    "employee relations",
    "offboarding",
    "classification",
    "exempt",
    "non-exempt",
    "offer letter",
    "onboarding",
]
STAGE_SIGNAL_PATTERNS = [
    re.compile(r"\bheadcount\b", flags=re.IGNORECASE),
    re.compile(r"\b\d{2,5}\s+(employees|employee|fte|team members)\b", flags=re.IGNORECASE),
    re.compile(r"\bseries\s+(a|b|c|d|e|f)\b", flags=re.IGNORECASE),
    re.compile(r"\bpre-ipo\b|\bipo\b|\bpublic company\b", flags=re.IGNORECASE),
]
IRRELEVANT_URL_TOKENS = [
    "press-release",
    "blog",
    "newsroom",
    "events",
    "webinar",
    "podcast",
    "investor",
    "marketing",
    "cookie",
    "privacy",
    "terms",
    "pricing",
    "sales",
    "install",
]
NON_US_LOCALE_TOKENS = [
    "india",
    "singapore",
    "australia",
    "germany",
    "france",
    "uk",
    "united-kingdom",
    "emea",
    "apac",
    "latam",
]
_DOMAIN_LAST_REQUEST_TS: Dict[str, float] = {}
_DOMAIN_LOCK = threading.Lock()
_INGESTION_STATS: Dict[str, object] = {}
logger = logging.getLogger(__name__)


def _env_int(name: str, default: int) -> int:
    value = os.getenv(name)
    if value is None:
        return default
    try:
        return int(value)
    except ValueError:
        return default


def _env_float(name: str, default: float) -> float:
    value = os.getenv(name)
    if value is None:
        return default
    try:
        return float(value)
    except ValueError:
        return default


def _env_bool(name: str, default: bool = False) -> bool:
    value = os.getenv(name)
    if value is None:
        return default
    return value.strip().lower() in {"1", "true", "yes", "y", "on"}


def _slugify(value: str) -> str:
    cleaned = re.sub(r"[^a-zA-Z0-9]+", "-", value.strip().lower()).strip("-")
    return cleaned or "document"


def _infer_file_title(path: Path, text: str) -> str:
    if path.suffix.lower() in {".md", ".markdown"}:
        for line in text.splitlines():
            stripped = line.strip()
            if not stripped:
                continue
            match = re.match(r"^\s*#\s+(.+?)\s*$", stripped)
            if match:
                title = " ".join(match.group(1).split()).strip()
                if title:
                    return title
                break
            # Stop scanning title candidates once body prose starts.
            if len(stripped.split()) >= 4:
                break
    return path.stem.replace("_", " ").replace("-", " ").strip() or "Untitled"


def _load_txt(path: Path) -> str:
    return normalize_text(path.read_text(encoding="utf-8", errors="ignore"))


def _load_pdf(path: Path) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise RuntimeError("pypdf is required to load PDF files") from exc

    reader = PdfReader(str(path))
    pages = []
    for page in reader.pages:
        pages.append(page.extract_text() or "")
    return normalize_text("\n\n".join(pages))


def _load_docx(path: Path) -> str:
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("python-docx is required to load DOCX files") from exc

    document = Document(str(path))
    return normalize_text("\n".join(p.text for p in document.paragraphs))


def _load_csv(path: Path) -> str:
    rows: List[List[str]] = []
    with path.open("r", encoding="utf-8", errors="ignore", newline="") as f:
        reader = csv.reader(f)
        for row in reader:
            rows.append([str(cell).strip() for cell in row])

    if not rows:
        return ""

    headers = rows[0]
    lines: List[str] = []
    if any(header.strip() for header in headers):
        lines.append("CSV headers: " + " | ".join(header.strip() or f"column_{idx + 1}" for idx, header in enumerate(headers)))

    for row in rows[1:]:
        parts: List[str] = []
        for idx, value in enumerate(row):
            cleaned = value.strip()
            if not cleaned:
                continue
            header = headers[idx].strip() if idx < len(headers) else f"column_{idx + 1}"
            header = header or f"column_{idx + 1}"
            parts.append(f"{header}: {cleaned}")
        if parts:
            lines.append(" | ".join(parts))

    return normalize_text("\n".join(lines))


def load_file_text(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix in {".txt", ".md"}:
        return _load_txt(path)
    if suffix == ".pdf":
        return _load_pdf(path)
    if suffix == ".docx":
        return _load_docx(path)
    if suffix == ".csv":
        return _load_csv(path)
    raise ValueError(f"Unsupported file type: {suffix}")


def _extract_readability_text(html: str) -> str:
    try:
        from readability import Document
    except ImportError:
        return ""

    try:
        summary_html = Document(html).summary()
    except Exception:
        return ""

    if not summary_html:
        return ""
    return normalize_text(_strip_html_fallback(summary_html))


def _strip_html_fallback(html: str) -> str:
    try:
        from bs4 import BeautifulSoup
    except ImportError as exc:
        raise RuntimeError("beautifulsoup4 is required to parse URL text") from exc

    soup = BeautifulSoup(html, "html.parser")
    for tag in soup(
        [
            "script",
            "style",
            "noscript",
            "nav",
            "header",
            "footer",
            "aside",
            "form",
            "svg",
        ]
    ):
        tag.extract()

    noisy_container_tokens = {
        "toc",
        "table-of-contents",
        "sidebar",
        "menu",
        "breadcrumb",
        "pagination",
        "footer",
        "header",
        "masthead",
        "drawer",
        "tabs",
    }
    for tag in soup.find_all(True):
        class_tokens = " ".join(tag.get("class", [])).lower()
        id_token = str(tag.get("id", "")).lower()
        bag = f"{class_tokens} {id_token}".strip()
        if any(token in bag for token in noisy_container_tokens):
            tag.extract()

    text = soup.get_text(separator="\n\n")
    lines = []
    for raw in text.splitlines():
        line = " ".join(raw.split())
        if not line:
            continue
        lines.append(line)
    return normalize_text("\n\n".join(lines))


def _extract_policy_wiki_text(html: str) -> str:
    try:
        from bs4 import BeautifulSoup
    except ImportError as exc:
        raise RuntimeError("beautifulsoup4 is required to parse URL text") from exc

    soup = BeautifulSoup(html, "html.parser")
    for tag in soup(["script", "style", "noscript", "nav", "footer", "aside", "form", "svg"]):
        tag.extract()

    lines: List[str] = []
    selectors = ["h1", "h2", "h3", "h4", "p", "li", "th", "td", "pre", "code"]
    for node in soup.find_all(selectors):
        text = " ".join(node.get_text(" ", strip=True).split())
        if not text:
            continue
        name = node.name.lower()
        if name in {"h1", "h2", "h3", "h4"}:
            lines.append(f"\n## {text}")
        elif name == "li":
            lines.append(f"- {text}")
        else:
            lines.append(text)
    return normalize_text("\n".join(lines))


def _is_policy_wiki_mode(url: str, html: str) -> bool:
    path = urlparse(url).path.lower()
    path_tokens = [
        "handbook",
        "wiki",
        "policy",
        "policies",
        "docs",
        "knowledge-base",
        "kb",
    ]
    if any(token in path for token in path_tokens):
        return True
    lower = html.lower()
    return any(token in lower for token in ["table of contents", "job families", "handbook"])


def _extract_main_content(html: str, url: str) -> Tuple[str, str]:
    if _is_policy_wiki_mode(url, html):
        policy_text = _extract_policy_wiki_text(html)
        if len(policy_text) >= 200:
            return policy_text, "policy_wiki"

    readable = _extract_readability_text(html)
    if len(readable) >= 400:
        return readable, "readability"
    return _strip_html_fallback(html), "fallback"


def _compute_link_ratio(html: str) -> float:
    try:
        from bs4 import BeautifulSoup
    except ImportError as exc:
        raise RuntimeError("beautifulsoup4 is required to parse URL text") from exc

    soup = BeautifulSoup(html, "html.parser")
    visible_text = " ".join(soup.stripped_strings)
    total_chars = max(len(visible_text), 1)
    link_chars = 0
    for anchor in soup.find_all("a"):
        anchor_text = " ".join(anchor.stripped_strings)
        link_chars += len(anchor_text)
    return link_chars / total_chars


def _avg_line_len(text: str) -> float:
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    if not lines:
        return 0.0
    return sum(len(line) for line in lines) / len(lines)


def _is_directory_page(text: str, link_ratio: float, avg_line_len: float) -> bool:
    if len(text) < 800:
        return True
    if link_ratio > 0.35:
        return True
    if avg_line_len < 55:
        return True
    return False


def _same_domain(url_a: str, url_b: str) -> bool:
    parsed_a = urlparse(url_a)
    parsed_b = urlparse(url_b)
    host_a = parsed_a.netloc.lower().removeprefix("www.")
    host_b = parsed_b.netloc.lower().removeprefix("www.")
    return host_a == host_b


def _extract_links(html: str, base_url: str) -> List[Tuple[str, str]]:
    try:
        from bs4 import BeautifulSoup
    except ImportError as exc:
        raise RuntimeError("beautifulsoup4 is required to parse URL text") from exc

    soup = BeautifulSoup(html, "html.parser")
    links: List[Tuple[str, str]] = []
    for anchor in soup.find_all("a", href=True):
        href = str(anchor.get("href", "")).strip()
        if not href:
            continue
        absolute = urljoin(base_url, href)
        absolute = urldefrag(absolute).url
        parsed = urlparse(absolute)
        if parsed.scheme not in {"http", "https"} or not parsed.netloc:
            continue
        anchor_text = " ".join(anchor.stripped_strings).strip()
        links.append((parsed.geturl(), anchor_text))
    return links


def _link_priority_score(link: str, anchor_text: str, locale_bias: str = "us") -> int:
    haystack = f"{link.lower()} {anchor_text.lower()}"
    score = sum(1 for token in DIRECTORY_LINK_KEYWORDS if token in haystack)
    score += 1 if any(token in haystack for token in ["policy", "handbook", "people"]) else 0

    # Down-rank obvious non-assessment pages and locale-mismatched paths.
    score -= sum(2 for token in IRRELEVANT_URL_TOKENS if token in haystack)
    if locale_bias.strip().lower() in {"us", "usa", "na"}:
        score -= sum(2 for token in NON_US_LOCALE_TOKENS if token in haystack)
    return score


def _url_page_relevance_score(*, url: str, title: str | None, text: str) -> int:
    lower_url = url.lower()
    path = urlparse(url).path.lower()
    title_lower = (title or "").lower()
    preview = " ".join(text.lower().split())
    if len(preview) > 3500:
        preview = preview[:3500]

    score = 0
    path_hits = sum(1 for token in HR_URL_PATH_TOKENS if token in path)
    score += min(5, path_hits * 2)
    score += min(3, sum(1 for token in DIRECTORY_LINK_KEYWORDS if token in title_lower))
    score += min(6, sum(1 for token in HR_CONTENT_TOKENS if token in preview))

    if any(pattern.search(preview) for pattern in STAGE_SIGNAL_PATTERNS):
        score += 3

    irrelevant_hits = sum(1 for token in IRRELEVANT_URL_TOKENS if token in lower_url)
    if irrelevant_hits:
        score -= min(6, irrelevant_hits * 2)

    if len(text.strip()) < 220:
        score -= 1
    return score


def _is_hr_or_stage_relevant_page(page: "_FetchedPage") -> bool:
    score = _url_page_relevance_score(url=page.final_url, title=page.title, text=page.text)
    if score >= 2:
        return True
    # Keep explicitly supplied seed URLs even if sparse; caller may have included strategic pages.
    if page.requested_url == page.final_url:
        return True
    return False


def _priority_directory_links(
    base_url: str,
    links: List[Tuple[str, str]],
    *,
    max_links: int = 25,
    locale_bias: str = "us",
    max_urls_per_domain: int = 8,
) -> List[str]:
    scored: List[Tuple[int, str]] = []
    seen = set()
    for link, anchor_text in links:
        if link in seen:
            continue
        seen.add(link)
        if not _same_domain(base_url, link):
            continue
        score = _link_priority_score(link, anchor_text, locale_bias=locale_bias)
        scored.append((score, link))

    scored.sort(key=lambda item: (-item[0], len(urlparse(item[1]).path), item[1]))
    selected: List[str] = []
    host_counts: Dict[str, int] = {}
    for score, link in scored:
        if score < 1:
            continue
        host = urlparse(link).netloc.lower().removeprefix("www.")
        current = host_counts.get(host, 0)
        if current >= max_urls_per_domain:
            continue
        host_counts[host] = current + 1
        selected.append(link)
        if len(selected) >= max_links:
            break
    return selected


def _normalized_paragraph_hash(paragraph: str) -> str:
    normalized = re.sub(r"\s+", " ", paragraph.strip().lower())
    normalized = re.sub(r"[^a-z0-9 ]+", "", normalized)
    return hashlib.sha1(normalized.encode("utf-8")).hexdigest()


def _split_paragraphs(text: str) -> List[str]:
    parts = [part.strip() for part in re.split(r"\n\s*\n+", text) if part.strip()]
    if parts:
        return parts
    return [line.strip() for line in text.splitlines() if line.strip()]


def _remove_repeated_boilerplate(pages: List["_FetchedPage"]) -> List["_FetchedPage"]:
    if len(pages) <= 1:
        return pages

    page_presence: Dict[str, int] = {}
    paragraph_hashes_per_page: List[set[str]] = []

    for page in pages:
        paragraphs = _split_paragraphs(page.text)
        hashes = {_normalized_paragraph_hash(p) for p in paragraphs if len(p) >= 40}
        paragraph_hashes_per_page.append(hashes)
        for hsh in hashes:
            page_presence[hsh] = page_presence.get(hsh, 0) + 1

    threshold = max(2, math.floor(len(pages) * 0.3) + 1)
    noisy_hashes = {hsh for hsh, count in page_presence.items() if count >= threshold}

    cleaned: List[_FetchedPage] = []
    for idx, page in enumerate(pages):
        paragraphs = _split_paragraphs(page.text)
        kept = []
        for paragraph in paragraphs:
            hsh = _normalized_paragraph_hash(paragraph)
            if hsh in noisy_hashes:
                continue
            kept.append(paragraph)
        cleaned_text = "\n\n".join(kept).strip()
        cleaned.append(replace(page, text=cleaned_text))
    return cleaned


@dataclass(frozen=True)
class _FetchedPage:
    requested_url: str
    final_url: str
    canonical_url: str
    text: str
    is_directory: bool
    links: List[Tuple[str, str]]
    status_code: int
    content_type: str
    retrieved_at: str
    title: str | None
    extract_mode: str


def _domain_throttle(url: str, delay_ms: int) -> None:
    if delay_ms <= 0:
        return
    domain = urlparse(url).netloc.lower()
    if not domain:
        return
    delay_s = delay_ms / 1000.0
    with _DOMAIN_LOCK:
        previous = _DOMAIN_LAST_REQUEST_TS.get(domain, 0.0)
        now = time.time()
        to_wait = delay_s - (now - previous)
        if to_wait > 0:
            time.sleep(to_wait)
        _DOMAIN_LAST_REQUEST_TS[domain] = time.time()


def _robots_allowed(url: str) -> bool:
    if not _env_bool("HR_REPORT_RESPECT_ROBOTS", default=False):
        return True
    try:
        from urllib.robotparser import RobotFileParser
    except Exception:
        return True
    parsed = urlparse(url)
    robots_url = f"{parsed.scheme}://{parsed.netloc}/robots.txt"
    rp = RobotFileParser()
    rp.set_url(robots_url)
    try:
        rp.read()
    except Exception:
        return True
    return bool(rp.can_fetch("HRReporterBot", url))


def _truncate_sentences(text: str, max_sentences: int = 3) -> str:
    sentences = [s.strip() for s in re.split(r"(?<=[.!?])\s+", text) if s.strip()]
    if not sentences:
        return text
    return " ".join(sentences[:max_sentences])


def _fetch_page(url: str, timeout_seconds: int = 15, light: bool = False) -> _FetchedPage:
    try:
        import requests
    except ImportError as exc:
        raise RuntimeError("requests is required to fetch URL text") from exc

    if not _robots_allowed(url):
        logger.info("robots_blocked url=%s", url)
        raise RuntimeError(f"Blocked by robots policy: {url}")

    retries = max(0, _env_int("HR_REPORT_URL_RETRIES", 2))
    backoff_base = max(0.1, _env_float("HR_REPORT_URL_BACKOFF_SECONDS", 0.6))
    throttle_ms = max(0, _env_int("HR_REPORT_URL_DOMAIN_THROTTLE_MS", 150))
    response = None
    last_exc: Exception | None = None
    for attempt in range(retries + 1):
        _domain_throttle(url, throttle_ms)
        try:
            response = requests.get(url, timeout=timeout_seconds, allow_redirects=True)
            if response.status_code in {429, 503}:
                logger.debug(
                    "transient_status_retry url=%s status=%s attempt=%s",
                    url,
                    response.status_code,
                    attempt,
                )
                time.sleep(backoff_base * (2**attempt))
                continue
            response.raise_for_status()
            break
        except Exception as exc:  # pragma: no cover - network edge
            last_exc = exc
            if attempt >= retries:
                raise
            time.sleep(backoff_base * (2**attempt))
    if response is None:
        if last_exc:
            raise last_exc
        raise RuntimeError(f"Failed to fetch URL: {url}")

    final_url = response.url or url
    canonical_url = final_url
    content_type = response.headers.get("content-type", "").lower()
    is_html = "text/html" in content_type or "<html" in response.text[:500].lower()
    retrieved_at = datetime.now(timezone.utc).isoformat()
    title = None
    extract_mode = "plain"

    if not is_html:
        text = normalize_text(response.text.strip())
        if light:
            text = _truncate_sentences(text, max_sentences=3)
        return _FetchedPage(
            requested_url=url,
            final_url=final_url,
            canonical_url=canonical_url,
            text=text,
            is_directory=False,
            links=[],
            status_code=response.status_code,
            content_type=content_type,
            retrieved_at=retrieved_at,
            title=title,
            extract_mode=extract_mode,
        )

    html = response.text
    try:
        from bs4 import BeautifulSoup
        soup = BeautifulSoup(html, "html.parser")
        title = " ".join((soup.title.get_text(" ", strip=True) if soup.title else "").split()) or None
    except Exception:
        title = None

    text, extract_mode = _extract_main_content(html, final_url)
    if light:
        text = _truncate_sentences(text, max_sentences=max(1, _env_int("HR_REPORT_URL_LIGHT_SENTENCES", 3)))
    link_ratio = _compute_link_ratio(html)
    avg_len = _avg_line_len(text)
    is_directory = _is_directory_page(text, link_ratio, avg_len)
    logger.debug(
        "fetched_page url=%s final_url=%s mode=%s directory=%s link_ratio=%.3f avg_len=%.1f",
        url,
        final_url,
        extract_mode,
        is_directory,
        link_ratio,
        avg_len,
    )
    links = _extract_links(html, final_url)
    return _FetchedPage(
        requested_url=url,
        final_url=final_url,
        canonical_url=canonical_url,
        text=text.strip(),
        is_directory=is_directory,
        links=links,
        status_code=response.status_code,
        content_type=content_type,
        retrieved_at=retrieved_at,
        title=title,
        extract_mode=extract_mode,
    )


def _fetch_url_set(
    url: str,
    timeout_seconds: int = 15,
    crawl_limit: int = 25,
    locale_bias: str = "us",
    max_urls_per_domain: int = 8,
) -> List[_FetchedPage]:
    return _fetch_url_set_with_limits(
        url=url,
        timeout_seconds=timeout_seconds,
        crawl_limit=crawl_limit,
        max_pages=None,
        child_workers=4,
        locale_bias=locale_bias,
        max_urls_per_domain=max_urls_per_domain,
    )


def _fetch_url_set_with_limits(
    *,
    url: str,
    timeout_seconds: int,
    crawl_limit: int,
    max_pages: int | None,
    child_workers: int,
    locale_bias: str = "us",
    max_urls_per_domain: int = 8,
) -> List[_FetchedPage]:
    root = _fetch_page(url, timeout_seconds=timeout_seconds, light=False)
    pages: List[_FetchedPage] = [root]
    request_budget = max_pages
    if request_budget is not None:
        request_budget = max(0, request_budget - 1)  # root request already consumed
    if not root.is_directory or crawl_limit <= 0:
        return pages

    visited = {root.final_url, root.requested_url}

    child_urls = _priority_directory_links(
        root.final_url,
        root.links,
        max_links=crawl_limit,
        locale_bias=locale_bias,
        max_urls_per_domain=max_urls_per_domain,
    )
    logger.debug(
        "directory_detected root=%s candidate_children=%s crawl_limit=%s",
        root.final_url,
        len(child_urls),
        crawl_limit,
    )
    if max_pages is not None:
        # Reserve one slot for the root page.
        allowed_children = max(0, max_pages - 1)
        child_urls = child_urls[:allowed_children]

    child_urls = [child_url for child_url in child_urls if child_url not in visited]
    visited.update(child_urls)

    if not child_urls:
        return pages

    # Phase 1: fetch light previews.
    preview_pages: List[_FetchedPage] = []
    workers = max(1, child_workers)
    if request_budget is not None:
        child_urls = child_urls[:request_budget]
    with ThreadPoolExecutor(max_workers=workers) as executor:
        future_map = {
            executor.submit(_fetch_page, child_url, timeout_seconds, True): child_url
            for child_url in child_urls
        }
        for future in as_completed(future_map):
            try:
                preview = future.result()
            except Exception:
                continue
            preview_pages.append(preview)
            if request_budget is not None:
                request_budget = max(0, request_budget - 1)

    if not preview_pages:
        return pages

    top_k = max(1, _env_int("HR_REPORT_URL_DIRECTORY_TOPK", 8))
    scored_previews = []
    for preview in preview_pages:
        score = _link_priority_score(preview.final_url, preview.text[:200], locale_bias=locale_bias)
        score += sum(1 for token in DIRECTORY_LINK_KEYWORDS if token in preview.text.lower())
        scored_previews.append((score, preview))
    scored_previews.sort(key=lambda item: item[0], reverse=True)
    selected_previews = [item[1] for item in scored_previews[:top_k]]
    logger.debug(
        "directory_selection root=%s preview_count=%s top_k=%s",
        root.final_url,
        len(preview_pages),
        top_k,
    )

    # Phase 2: fetch full content for selected candidates.
    selected_urls = [p.final_url for p in selected_previews]
    if request_budget is not None and request_budget <= 0:
        pages.extend(selected_previews)
        return pages
    if request_budget is not None:
        selected_urls = selected_urls[:request_budget]
    with ThreadPoolExecutor(max_workers=workers) as executor:
        future_map = {
            executor.submit(_fetch_page, child_url, timeout_seconds, False): child_url
            for child_url in selected_urls
        }
        for future in as_completed(future_map):
            if max_pages is not None and len(pages) >= max_pages:
                break
            try:
                child = future.result()
            except Exception:
                continue
            pages.append(child)
            if request_budget is not None:
                request_budget = max(0, request_budget - 1)
    return pages


def parse_url_candidates(raw_urls: Sequence[str] | str | None) -> List[str]:
    if not raw_urls:
        return []

    if isinstance(raw_urls, str):
        merged = raw_urls
    else:
        merged = "\n".join(raw_urls)

    # Normalize common malformed scheme copy/paste cases: https:/foo -> https://foo
    merged = re.sub(r"(?<!/)https:/(?=[^/])", "https://", merged, flags=re.IGNORECASE)
    merged = re.sub(r"(?<!/)http:/(?=[^/])", "http://", merged, flags=re.IGNORECASE)

    # Support concatenated URLs by splitting at the next scheme token.
    pattern = re.compile(r"https?://.*?(?=(?:https?://)|[\s,]|$)", flags=re.IGNORECASE)
    candidates = [match.group(0).strip(" \t\r\n,;|)>]}\"'") for match in pattern.finditer(merged)]

    deduped: List[str] = []
    seen = set()
    for candidate in candidates:
        parsed = urlparse(candidate)
        if parsed.scheme not in {"http", "https"} or not parsed.netloc:
            continue
        normalized = parsed.geturl()
        if normalized in seen:
            continue
        seen.add(normalized)
        deduped.append(normalized)

    return deduped


def _discover_input_files(input_path: Path) -> List[Path]:
    if not input_path.exists():
        raise FileNotFoundError(f"Input path does not exist: {input_path}")

    if input_path.is_file():
        return [input_path]

    files: List[Path] = []
    for path in sorted(input_path.rglob("*")):
        if path.is_file() and path.suffix.lower() in SUPPORTED_EXTENSIONS:
            files.append(path)
    return files


def _relative_file_source(file_path: Path, input_root: Path | None) -> str:
    resolved = file_path.resolve()
    if input_root is not None:
        try:
            return resolved.relative_to(input_root.resolve()).as_posix()
        except ValueError:
            pass
    return file_path.name


def load_documents(
    input_path: str | Path | None = None,
    pasted_text: str | None = None,
    urls: Sequence[str] | None = None,
    max_urls_per_domain: int | None = None,
    max_total_urls: int | None = None,
    locale_bias: str | None = None,
) -> List[RawDocument]:
    docs: List[RawDocument] = []
    counter = 1
    url_timeout_seconds = max(1, _env_int("HR_REPORT_URL_TIMEOUT_SECONDS", 8))
    crawl_limit = max(0, max_urls_per_domain or _env_int("HR_REPORT_URL_CRAWL_LIMIT", 25))
    crawl_max_seed_urls = max(0, _env_int("HR_REPORT_URL_CRAWL_MAX_SEED_URLS", 8))
    url_max_total_fetches = max(1, max_total_urls or _env_int("HR_REPORT_URL_MAX_TOTAL_FETCHES", 120))
    url_child_workers = max(1, _env_int("HR_REPORT_URL_CHILD_WORKERS", 6))
    locale = (locale_bias or os.getenv("HR_REPORT_URL_LOCALE_BIAS", "us")).strip().lower() or "us"

    stats = {
        "urls_seeded": 0,
        "urls_fetched": 0,
        "url_fetch_failures": 0,
        "url_redirects": 0,
        "url_pages_filtered_low_relevance": 0,
        "crawl_enabled": False,
        "crawl_depth": 1,
        "directory_crawl_ran": False,
        "domains": [],
    }
    logger.info(
        "ingestion_start input_path=%s pasted_text=%s urls=%s",
        str(input_path) if input_path else "none",
        bool(pasted_text and pasted_text.strip()),
        len(parse_url_candidates(urls)),
    )

    input_root: Path | None = None
    if input_path:
        base = Path(input_path)
        input_root = base if base.is_dir() else base.parent
        for file_path in _discover_input_files(base):
            text = load_file_text(file_path)
            if not text.strip():
                continue
            doc_id = f"doc-{counter:03d}-{_slugify(file_path.stem)}"
            counter += 1
            relative_source = _relative_file_source(file_path, input_root)
            docs.append(
                RawDocument(
                    doc_id=doc_id,
                    source_id=f"file:{relative_source}",
                    source=relative_source,
                    source_type="file",
                    title=_infer_file_title(file_path, text),
                    retrieved_at=datetime.now(timezone.utc).isoformat(),
                    content_hash=hashlib.sha1(text.encode("utf-8")).hexdigest(),
                    metadata={"absolute_path": str(file_path.resolve())},
                    text=text,
                )
            )
            logger.debug("ingestion_file doc_id=%s path=%s chars=%s", doc_id, file_path, len(text))

    if pasted_text and pasted_text.strip():
        docs.append(
            RawDocument(
                doc_id=f"doc-{counter:03d}-pasted-text",
                source_id="text:pasted",
                source="Pasted input text",
                source_type="text",
                retrieved_at=datetime.now(timezone.utc).isoformat(),
                content_hash=hashlib.sha1(normalize_text(pasted_text).encode("utf-8")).hexdigest(),
                text=normalize_text(pasted_text),
            )
        )
        counter += 1
        logger.debug("ingestion_text doc_id=%s chars=%s", f"doc-{counter - 1:03d}-pasted-text", len(pasted_text))

    url_errors: List[str] = []
    fetched_pages: List[_FetchedPage] = []
    seed_urls = parse_url_candidates(urls)
    stats["urls_seeded"] = len(seed_urls)
    crawl_enabled = len(seed_urls) <= crawl_max_seed_urls
    stats["crawl_enabled"] = crawl_enabled
    remaining_fetch_budget = url_max_total_fetches
    logger.info(
        "ingestion_urls seeds=%s crawl_enabled=%s crawl_limit=%s max_fetches=%s",
        len(seed_urls),
        crawl_enabled,
        crawl_limit,
        url_max_total_fetches,
    )

    for url in seed_urls:
        if remaining_fetch_budget <= 0:
            break

        effective_crawl_limit = crawl_limit if crawl_enabled else 0
        try:
            pages = _fetch_url_set_with_limits(
                url=url,
                timeout_seconds=url_timeout_seconds,
                crawl_limit=effective_crawl_limit,
                max_pages=remaining_fetch_budget,
                child_workers=url_child_workers,
                locale_bias=locale,
                max_urls_per_domain=max(1, crawl_limit),
            )
            fetched_pages.extend(pages)
            stats["urls_fetched"] = int(stats["urls_fetched"]) + len(pages)
            stats["directory_crawl_ran"] = bool(stats["directory_crawl_ran"]) or (
                len(pages) > 1
            )
            logger.info(
                "url_seed_processed url=%s pages=%s crawl_enabled=%s",
                url,
                len(pages),
                crawl_enabled,
            )
            remaining_fetch_budget -= len(pages)
        except Exception as exc:
            stats["url_fetch_failures"] = int(stats["url_fetch_failures"]) + 1
            url_errors.append(f"{url}: {exc}")
            logger.warning("url_seed_failed url=%s error=%s", url, exc)
    fetched_pages = _remove_repeated_boilerplate(fetched_pages)

    seen_sources = set()
    domains = set()
    for page in fetched_pages:
        if not page.text.strip():
            continue
        if not _is_hr_or_stage_relevant_page(page):
            stats["url_pages_filtered_low_relevance"] = int(stats["url_pages_filtered_low_relevance"]) + 1
            logger.debug("url_page_skipped_low_relevance url=%s title=%s", page.final_url, page.title or "")
            continue
        if page.final_url in seen_sources:
            continue
        seen_sources.add(page.final_url)
        parsed = urlparse(page.final_url)
        if parsed.netloc:
            domains.add(parsed.netloc.lower())
        if page.requested_url != page.final_url:
            stats["url_redirects"] = int(stats["url_redirects"]) + 1
        docs.append(
            RawDocument(
                doc_id=f"doc-{counter:03d}-{_slugify(page.final_url)}",
                source_id=f"url:{page.final_url}",
                source=page.final_url,
                source_type="url",
                canonical_url=page.canonical_url,
                final_url=page.final_url,
                retrieved_at=page.retrieved_at,
                title=page.title,
                content_hash=hashlib.sha1(page.text.encode("utf-8")).hexdigest(),
                http_status=page.status_code,
                content_type=page.content_type,
                metadata={
                    "extract_mode": page.extract_mode,
                    "is_directory": page.is_directory,
                },
                text=normalize_text(page.text),
            )
        )
        counter += 1

    stats["domains"] = sorted(domains)
    global _INGESTION_STATS
    _INGESTION_STATS = stats
    logger.info(
        "ingestion_complete docs=%s urls_seeded=%s urls_fetched=%s failures=%s",
        len(docs),
        stats["urls_seeded"],
        stats["urls_fetched"],
        stats["url_fetch_failures"],
    )

    if not docs:
        message = "No input text found. Provide --input, --text, or --url."
        if url_errors:
            message += f" URL fetch errors: {' | '.join(url_errors)}"
        raise ValueError(message)

    return docs


def get_last_ingestion_stats() -> Dict[str, object]:
    return dict(_INGESTION_STATS)
