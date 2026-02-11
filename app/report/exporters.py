from __future__ import annotations

from io import BytesIO
import re
import textwrap
from typing import List, Tuple

REPORT_TITLE = "HR Assessment Report"


def _strip_inline_markdown(text: str) -> str:
    text = text.replace("**", "").replace("__", "")
    text = text.replace("`", "")
    text = re.sub(r"\[(.*?)\]\((.*?)\)", r"\1 (\2)", text)
    return text.strip()


def _markdown_to_semantic_lines(markdown: str) -> List[Tuple[str, str]]:
    lines: List[Tuple[str, str]] = []
    for raw in markdown.splitlines():
        line = raw.rstrip()
        stripped = line.strip()
        if not stripped:
            lines.append(("blank", ""))
            continue

        if re.match(r"^\|\s*-", stripped):
            # Markdown table separator row.
            continue

        if stripped.startswith("### "):
            lines.append(("h3", _strip_inline_markdown(stripped[4:])))
            continue
        if stripped.startswith("## "):
            lines.append(("h2", _strip_inline_markdown(stripped[3:])))
            continue
        if stripped.startswith("# "):
            lines.append(("h1", _strip_inline_markdown(stripped[2:])))
            continue
        if stripped.startswith("- "):
            lines.append(("bullet", _strip_inline_markdown(stripped[2:])))
            continue
        if re.match(r"^\d+\.\s+", stripped):
            lines.append(("number", _strip_inline_markdown(re.sub(r"^\d+\.\s+", "", stripped))))
            continue
        if stripped.startswith("|") and stripped.endswith("|"):
            columns = [item.strip() for item in stripped.strip("|").split("|")]
            lines.append(("table", " | ".join(_strip_inline_markdown(item) for item in columns)))
            continue

        lines.append(("paragraph", _strip_inline_markdown(stripped)))

    return lines


def markdown_to_docx_bytes(markdown: str, title: str = REPORT_TITLE) -> bytes:
    try:
        from docx import Document
    except ImportError as exc:  # pragma: no cover
        raise RuntimeError("python-docx is required for DOCX export") from exc

    document = Document()
    document.add_heading(title, level=0)

    for kind, content in _markdown_to_semantic_lines(markdown):
        if kind == "blank":
            document.add_paragraph("")
            continue
        if kind == "h1":
            document.add_heading(content, level=1)
            continue
        if kind == "h2":
            document.add_heading(content, level=2)
            continue
        if kind == "h3":
            document.add_heading(content, level=3)
            continue
        if kind == "bullet":
            document.add_paragraph(content, style="List Bullet")
            continue
        if kind == "number":
            document.add_paragraph(content, style="List Number")
            continue
        if kind == "table":
            document.add_paragraph(content)
            continue
        document.add_paragraph(content)

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _escape_pdf_text(text: str) -> str:
    return text.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")


def _build_pdf_pages(lines: List[str], *, max_chars: int = 95, max_lines_per_page: int = 48) -> List[List[str]]:
    wrapped: List[str] = []
    for line in lines:
        clean = line.strip()
        if not clean:
            wrapped.append("")
            continue
        wrapped.extend(textwrap.wrap(clean, width=max_chars) or [""])

    pages: List[List[str]] = []
    cursor = 0
    while cursor < len(wrapped):
        pages.append(wrapped[cursor : cursor + max_lines_per_page])
        cursor += max_lines_per_page
    if not pages:
        pages = [[""]]
    return pages


def _page_stream(lines: List[str], *, margin_x: int = 48, top_y: int = 760, font_size: int = 10, leading: int = 14) -> bytes:
    cmds: List[str] = [
        "BT",
        f"/F1 {font_size} Tf",
        f"{leading} TL",
        f"{margin_x} {top_y} Td",
    ]
    for idx, line in enumerate(lines):
        escaped = _escape_pdf_text(line)
        if idx > 0:
            cmds.append("T*")
        cmds.append(f"({escaped}) Tj")
    cmds.append("ET")
    return "\n".join(cmds).encode("latin-1", errors="replace")


def markdown_to_pdf_bytes(markdown: str, title: str = REPORT_TITLE) -> bytes:
    semantic_lines = _markdown_to_semantic_lines(markdown)

    printable: List[str] = [title, ""]
    for kind, content in semantic_lines:
        if kind == "blank":
            printable.append("")
        elif kind == "h1":
            printable.extend([content.upper(), ""])
        elif kind == "h2":
            printable.extend([content, ""])
        elif kind == "h3":
            printable.append(content)
        elif kind == "bullet":
            printable.append(f"- {content}")
        elif kind == "number":
            printable.append(f"- {content}")
        elif kind == "table":
            printable.append(content)
        else:
            printable.append(content)

    pages = _build_pdf_pages(printable)
    page_streams = [_page_stream(page_lines) for page_lines in pages]

    # Object layout:
    # 1: catalog
    # 2: pages tree
    # 3: font
    # 4..: alternating page and content objects
    objects: List[bytes] = []

    objects.append(b"<< /Type /Catalog /Pages 2 0 R >>")

    page_object_ids: List[int] = []
    content_object_ids: List[int] = []
    next_id = 4
    for _ in page_streams:
        page_object_ids.append(next_id)
        content_object_ids.append(next_id + 1)
        next_id += 2

    kids = " ".join(f"{obj_id} 0 R" for obj_id in page_object_ids)
    objects.append(f"<< /Type /Pages /Kids [{kids}] /Count {len(page_object_ids)} >>".encode("latin-1"))

    objects.append(b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>")

    for page_obj_id, content_obj_id, stream in zip(page_object_ids, content_object_ids, page_streams):
        page_obj = (
            "<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
            f"/Resources << /Font << /F1 3 0 R >> >> /Contents {content_obj_id} 0 R >>"
        )
        objects.append(page_obj.encode("latin-1"))

        content_header = f"<< /Length {len(stream)} >>\nstream\n".encode("latin-1")
        content_footer = b"\nendstream"
        objects.append(content_header + stream + content_footer)

    pdf = BytesIO()
    pdf.write(b"%PDF-1.4\n%\xe2\xe3\xcf\xd3\n")

    offsets: List[int] = [0]
    for idx, body in enumerate(objects, start=1):
        offsets.append(pdf.tell())
        pdf.write(f"{idx} 0 obj\n".encode("latin-1"))
        pdf.write(body)
        pdf.write(b"\nendobj\n")

    xref_start = pdf.tell()
    pdf.write(f"xref\n0 {len(offsets)}\n".encode("latin-1"))
    pdf.write(b"0000000000 65535 f \n")
    for offset in offsets[1:]:
        pdf.write(f"{offset:010d} 00000 n \n".encode("latin-1"))

    trailer = (
        f"trailer\n<< /Size {len(offsets)} /Root 1 0 R >>\n"
        f"startxref\n{xref_start}\n%%EOF\n"
    )
    pdf.write(trailer.encode("latin-1"))
    return pdf.getvalue()
